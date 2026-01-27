"""
Build a single docx from docs/project/engineering_book/* using a manifest.

Markdown is the source of truth. Docx is the shareable artifact.
"""

from __future__ import annotations

import argparse
from pathlib import Path
import re

from docx import Document
from docx.shared import Pt


_RE_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_RE_ORDERED = re.compile(r"^\s*(\d+)[\.\)]\s+(.*)$")
_RE_BULLET = re.compile(r"^\s*-\s+(.*)$")


def _add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph()
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def _append_md(doc: Document, md_text: str) -> None:
    in_code = False
    code_lines: list[str] = []

    for raw in md_text.splitlines():
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            if in_code:
                _add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        m = _RE_HEADING.match(line)
        if m:
            level = len(m.group(1))
            title = m.group(2).strip()
            doc.add_paragraph(title, style=f"Heading {min(level, 4)}")
            continue

        m = _RE_ORDERED.match(line)
        if m:
            doc.add_paragraph(m.group(2).strip(), style="List Number")
            continue

        m = _RE_BULLET.match(line)
        if m:
            doc.add_paragraph(m.group(1).strip(), style="List Bullet")
            continue

        doc.add_paragraph(line)

    if in_code and code_lines:
        _add_code_block(doc, code_lines)


def _load_manifest(path: Path) -> list[str]:
    out = []
    for raw in path.read_text(encoding="utf-8").splitlines():
        s = raw.strip()
        if not s or s.startswith("#"):
            continue
        out.append(s)
    return out


def build(manifest: Path, out_path: Path) -> None:
    root = manifest.parent.parent  # docs/project/engineering_book
    doc = Document()

    for rel in _load_manifest(manifest):
        md_path = (root / rel).resolve()
        if not md_path.exists():
            raise FileNotFoundError(f"Missing in manifest: {rel} -> {md_path}")
        _append_md(doc, md_path.read_text(encoding="utf-8"))
        doc.add_page_break()

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--manifest", default="docs/project/engineering_book/_build/manifest.txt")
    parser.add_argument(
        "--out",
        default="docs/project/engineering_book/_build/NSGABlack_Engineering_Book.docx",
    )
    args = parser.parse_args()
    build(Path(args.manifest).resolve(), Path(args.out).resolve())
    print(f"[docx] wrote: {Path(args.out).resolve()}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

