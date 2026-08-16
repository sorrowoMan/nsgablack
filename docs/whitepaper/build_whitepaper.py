"""Build the Chinese Black Framework Stack whitepaper as Markdown and DOCX.

The builder intentionally reads the live blackbase/nsgablack/mlblack worktrees.
It does not modify source documents and never invokes Git cleanup commands.

DOCX design system:
- preset: compact_reference_guide
- first-page pattern: editorial_cover
- named override: Microsoft YaHei for East Asian glyphs
- named override: no decorative header rule on the cover
"""

from __future__ import annotations

import json
import re
from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


WHITEPAPER_DIR = Path(__file__).resolve().parent
NSGA_ROOT = WHITEPAPER_DIR.parents[1]
DESKTOP_ROOT = NSGA_ROOT.parent
REPO_ROOTS = {
    "nsgablack": NSGA_ROOT,
    "blackbase": DESKTOP_ROOT / "blackbase",
    "mlblack": DESKTOP_ROOT / "mlblack",
}

MANIFEST_PATH = WHITEPAPER_DIR / "manifest.json"
MARKDOWN_OUT = WHITEPAPER_DIR / "Black_Framework_Stack_Whitepaper_CN.md"
DOCX_OUT = WHITEPAPER_DIR / "Black_Framework_Stack_Whitepaper_CN.docx"
STATS_OUT = WHITEPAPER_DIR / "Black_Framework_Stack_Whitepaper_CN.stats.json"

# compact_reference_guide tokens
PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_TOP_BOTTOM_DXA = 80
CELL_MARGIN_START_END_DXA = 120
ASCII_FONT = "Calibri"
CJK_FONT = "Microsoft YaHei"
MONO_FONT = "Consolas"
HEADING_BLUE = "2E74B5"
HEADING_DARK_BLUE = "1F4D78"
INK_BLUE = "0B2545"
MUTED = "667085"
LIGHT_BLUE_GRAY = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT_FILL = "F4F6F9"
WHITE = "FFFFFF"


def _load_manifest() -> dict:
    return json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))


def _chapter_path(chapter: dict) -> Path:
    if chapter["kind"] == "original":
        return WHITEPAPER_DIR / chapter["path"]
    return REPO_ROOTS[chapter["repo"]] / chapter["path"]


def _strip_front_matter(text: str) -> str:
    if text.startswith("---\n"):
        end = text.find("\n---\n", 4)
        if end >= 0:
            return text[end + 5 :]
    return text


def _strip_first_h1(text: str) -> str:
    lines = text.splitlines()
    for idx, line in enumerate(lines):
        if line.startswith("# "):
            del lines[idx]
            break
        if line.strip():
            break
    return "\n".join(lines).strip()


def _source_note(chapter: dict) -> str:
    status = chapter.get("status", "说明")
    if chapter["kind"] == "original":
        return f"> 文档状态：{status}。本章为白皮书原创主卷。"
    return (
        f"> 文档状态：{status}；来源：`{chapter['repo']}/{chapter['path']}`。"
        "本节按当前工作树合订；如与原创主卷或实时源码冲突，以原创主卷标注的规范和实时源码为准。"
    )


def compose_markdown(manifest: dict) -> tuple[str, list[dict]]:
    resolved: list[dict] = []
    for index, chapter in enumerate(manifest["chapters"], start=1):
        path = _chapter_path(chapter)
        if not path.exists():
            raise FileNotFoundError(f"whitepaper source missing: {path}")
        source = _strip_front_matter(path.read_text(encoding="utf-8-sig"))
        body = _strip_first_h1(source)
        resolved.append({**chapter, "index": index, "path_obj": path, "body": body})

    lines = [
        f"# {manifest['title']}",
        "",
        manifest["subtitle"],
        "",
        f"版本：{manifest['edition']}  ",
        f"生成日期：{date.today().isoformat()}  ",
        "主版本：中文",
        "",
        "## 阅读目录",
        "",
    ]
    for chapter in resolved:
        lines.append(f"{chapter['index']}. {chapter['title']}（{chapter.get('status', '说明')}）")
    lines.extend(["", "---", ""])
    for chapter in resolved:
        lines.extend(
            [
                f"# {chapter['title']}",
                "",
                _source_note(chapter),
                "",
                chapter["body"],
                "",
                "---",
                "",
            ]
        )
    return "\n".join(lines).strip() + "\n", resolved


def _strip_markdown_for_count(text: str) -> str:
    value = re.sub(r"```[^\n]*\n", "", text)
    value = value.replace("```", "")
    value = re.sub(r"!\[([^]]*)\]\([^)]*\)", r"\1", value)
    value = re.sub(r"\[([^]]+)\]\([^)]*\)", r"\1", value)
    value = re.sub(r"[*_`>#|~-]", "", value)
    return value


def _stats(markdown: str, resolved: Sequence[dict], manifest: dict) -> dict:
    plain = _strip_markdown_for_count(markdown)
    return {
        "title": manifest["title"],
        "edition": manifest["edition"],
        "generated_on": date.today().isoformat(),
        "chapter_count": len(resolved),
        "original_chapter_count": sum(c["kind"] == "original" for c in resolved),
        "source_chapter_count": sum(c["kind"] == "source" for c in resolved),
        "raw_markdown_characters": len(markdown),
        "readable_characters_excluding_whitespace": len(re.sub(r"\s+", "", plain)),
        "cjk_characters": len(re.findall(r"[\u3400-\u9fff]", plain)),
        "target_characters": int(manifest.get("target_characters", 100000)),
        "source_paths": [str(c["path_obj"]) for c in resolved],
    }


def _set_run_font(
    run,
    *,
    ascii_font: str = ASCII_FONT,
    east_asia_font: str = CJK_FONT,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _set_style_font(style, *, size: float, color: str = "000000", bold: bool = False) -> None:
    font = style.font
    font.name = ASCII_FONT
    font.size = Pt(size)
    font.color.rgb = RGBColor.from_string(color)
    font.bold = bold
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), ASCII_FONT)
    rfonts.set(qn("w:hAnsi"), ASCII_FONT)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)


def _set_paragraph_spacing(style, *, before: float, after: float, line: float) -> None:
    fmt = style.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def _shade_paragraph(paragraph, fill: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_paragraph_left_border(paragraph, *, color: str, size: int = 12, space: int = 8) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    borders = ppr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        ppr.append(borders)
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)


def _configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    _set_style_font(normal, size=11)
    _set_paragraph_spacing(normal, before=0, after=6, line=1.25)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    heading_tokens = {
        "Heading 1": (16, HEADING_BLUE, 18, 10),
        "Heading 2": (13, HEADING_BLUE, 14, 7),
        "Heading 3": (12, HEADING_DARK_BLUE, 10, 5),
        "Heading 4": (11, INK_BLUE, 8, 4),
        "Heading 5": (10.5, INK_BLUE, 6, 3),
        "Heading 6": (10.5, MUTED, 6, 3),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        _set_style_font(style, size=size, color=color, bold=True)
        _set_paragraph_spacing(style, before=before, after=after, line=1.1)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for name in ("List Bullet", "List Bullet 2", "List Bullet 3", "List Number", "List Number 2", "List Number 3"):
        style = doc.styles[name]
        _set_style_font(style, size=11)
        _set_paragraph_spacing(style, before=0, after=4, line=1.25)
        level = 0
        if name.endswith(" 2"):
            level = 1
        elif name.endswith(" 3"):
            level = 2
        style.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
        style.paragraph_format.first_line_indent = Inches(-0.188)

    custom = {
        "WP Cover Kicker": (10.5, HEADING_BLUE, True, 0, 12, 1.0),
        "WP Cover Title": (30, INK_BLUE, True, 0, 8, 1.0),
        "WP Cover Subtitle": (15, HEADING_DARK_BLUE, False, 0, 20, 1.15),
        "WP Source Note": (9, MUTED, False, 0, 6, 1.15),
        "WP Quote": (10.5, INK_BLUE, False, 4, 6, 1.2),
        "WP Code": (8.5, "1F2937", False, 0, 0, 1.0),
        "WP Table Text": (8.5, "1F2937", False, 0, 0, 1.05),
    }
    for name, (size, color, bold, before, after, line) in custom.items():
        style = doc.styles[name] if name in doc.styles else doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        _set_style_font(style, size=size, color=color, bold=bold)
        _set_paragraph_spacing(style, before=before, after=after, line=line)


def _configure_page(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True


def _add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    _set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    field_run = OxmlElement("w:r")
    field_run.append(begin)
    field_run.append(instr)
    field_run.append(separate)
    field_run.append(value)
    field_run.append(end)
    paragraph._p.append(field_run)
    tail = paragraph.add_run(" 页")
    _set_run_font(tail, size=9, color=MUTED)


def _configure_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Black Framework Stack 中文白皮书")
    _set_run_font(run, size=9, color=MUTED, bold=True)
    footer = section.footer
    fp = footer.paragraphs[0]
    _add_page_number(fp)


def _next_numbering_id(numbering, tag: str, attr: str) -> int:
    values = []
    for node in numbering.findall(qn(tag)):
        raw = node.get(qn(attr))
        if raw is not None:
            values.append(int(raw))
    return max(values, default=0) + 1


def _create_decimal_numbering(doc: Document, *, start: int = 1) -> int:
    """Create a real, restartable three-level decimal numbering definition."""
    numbering = doc.part.numbering_part.element
    abstract_id = _next_numbering_id(numbering, "w:abstractNum", "w:abstractNumId")
    num_id = _next_numbering_id(numbering, "w:num", "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)
    for level in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start_node = OxmlElement("w:start")
        start_node.set(qn("w:val"), "1")
        lvl.append(start_node)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), f"%{level + 1}.")
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(540 + level * 360))
        tabs.append(tab)
        ppr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(540 + level * 360))
        ind.set(qn("w:hanging"), "270")
        ppr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.append(spacing)
        lvl.append(ppr)
        abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    if start != 1:
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), str(start))
        override.append(start_override)
        num.append(override)
    numbering.append(num)
    return num_id


def _apply_numbering(paragraph, *, num_id: int, level: int = 0) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.insert(0, num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def _add_cover(doc: Document, manifest: dict, stats: dict) -> None:
    for _ in range(5):
        doc.add_paragraph()
    kicker = doc.add_paragraph(style="WP Cover Kicker")
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.add_run("统一运行底座 · 优化搜索 · 机器学习语义")
    title = doc.add_paragraph(style="WP Cover Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(manifest["title"])
    subtitle = doc.add_paragraph(style="WP Cover Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(manifest["subtitle"])
    for _ in range(4):
        doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(5)
    r = meta.add_run(manifest["edition"])
    _set_run_font(r, size=11, color=INK_BLUE, bold=True)
    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta2.paragraph_format.space_after = Pt(4)
    r = meta2.add_run(
        f"生成日期 {stats['generated_on']} · {stats['chapter_count']} 章 · "
        f"可读字符 {stats['readable_characters_excluding_whitespace']:,}"
    )
    _set_run_font(r, size=9.5, color=MUTED)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(18)
    r = note.add_run("中文主版本。代码标识、协议名与命令保留原始英文。")
    _set_run_font(r, size=9, color=MUTED, italic=True)
    doc.add_page_break()


def _add_contents(doc: Document, resolved: Sequence[dict]) -> None:
    p = doc.add_paragraph("阅读目录", style="Heading 1")
    p.paragraph_format.page_break_before = False
    intro = doc.add_paragraph(
        "本目录按阅读顺序编排。前六章是统一叙事与运行规范，后续章节合订当前三仓的正式教程和参考材料。Word 导航窗格可按标题层级跳转。"
    )
    intro.paragraph_format.space_after = Pt(10)
    toc_num_id = _create_decimal_numbering(doc, start=1)
    for chapter in resolved:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        _apply_numbering(p, num_id=toc_num_id)
        p.add_run(f"{chapter['title']}  ")
        tag = p.add_run(f"[{chapter.get('status', '说明')}]")
        _set_run_font(tag, size=9, color=MUTED, bold=True)
    doc.add_page_break()


INLINE_TOKEN_RE = re.compile(
    r"(`[^`]+`|\*\*[^*]+\*\*|__[^_]+__|(?<!\*)\*[^*]+\*(?!\*)|(?<!_)_[^_]+_(?!_)|\[[^]]+\]\([^)]+\))"
)


def _add_inline(paragraph, text: str) -> None:
    text = re.sub(r"<br\s*/?>", " ", text, flags=re.I)
    text = re.sub(r"</?[^>]+>", "", text)
    pos = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            _set_run_font(run, size=None)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, ascii_font=MONO_FONT, east_asia_font=CJK_FONT, size=9.5, color=HEADING_DARK_BLUE)
        elif token.startswith("**") or token.startswith("__"):
            run = paragraph.add_run(token[2:-2])
            _set_run_font(run, bold=True)
        elif token.startswith("*") or token.startswith("_"):
            run = paragraph.add_run(token[1:-1])
            _set_run_font(run, italic=True)
        elif token.startswith("["):
            lm = re.match(r"\[([^]]+)\]\(([^)]+)\)", token)
            assert lm
            label, target = lm.groups()
            display = label if target.startswith((".", "#", "/")) else f"{label}（{target}）"
            run = paragraph.add_run(display)
            _set_run_font(run, color=HEADING_BLUE)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        _set_run_font(run, size=None)


def _is_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def _split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _set_cell_margins(cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", CELL_MARGIN_TOP_BOTTOM_DXA),
        ("bottom", CELL_MARGIN_TOP_BOTTOM_DXA),
        ("start", CELL_MARGIN_START_END_DXA),
        ("end", CELL_MARGIN_START_END_DXA),
    ):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: Sequence[int]) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            _set_cell_width(cell, width)
            _set_cell_margins(cell)


def _column_widths(rows: Sequence[Sequence[str]]) -> list[int]:
    count = max(len(row) for row in rows)
    weights: list[float] = []
    for idx in range(count):
        lengths = [len(re.sub(r"[`*_]", "", row[idx])) if idx < len(row) else 0 for row in rows]
        weights.append(float(max(6, min(42, max(lengths, default=6)))))
    total = sum(weights)
    widths = [max(700, round(CONTENT_WIDTH_DXA * value / total)) for value in weights]
    diff = CONTENT_WIDTH_DXA - sum(widths)
    widths[-1] += diff
    if widths[-1] < 600:
        need = 600 - widths[-1]
        widths[-1] = 600
        widths[widths.index(max(widths[:-1]))] -= need
    return widths


def _add_table(doc: Document, rows: Sequence[Sequence[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    normalized = [list(row) + [""] * (columns - len(row)) for row in rows]
    table = doc.add_table(rows=len(normalized), cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"
    widths = _column_widths(normalized)
    _set_table_geometry(table, widths)
    for ridx, values in enumerate(normalized):
        row = table.rows[ridx]
        if ridx == 0:
            tr_pr = row._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
        for cidx, value in enumerate(values):
            cell = row.cells[cidx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.style = doc.styles["WP Table Text"]
            p.paragraph_format.space_after = Pt(0)
            if ridx == 0:
                _shade_paragraph(p, LIGHT_BLUE_GRAY)
            _add_inline(p, value)
            for run in p.runs:
                _set_run_font(run, size=8.5, bold=True if ridx == 0 else None)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(3)


def _add_code_block(doc: Document, lines: Sequence[str], language: str = "") -> None:
    p = doc.add_paragraph(style="WP Code")
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    _shade_paragraph(p, LIGHT_GRAY)
    if language:
        label = p.add_run(f"{language}\n")
        _set_run_font(label, ascii_font=MONO_FONT, east_asia_font=CJK_FONT, size=7.5, color=MUTED, bold=True)
    content = "\n".join(lines).rstrip()
    run = p.add_run(content or " ")
    _set_run_font(run, ascii_font=MONO_FONT, east_asia_font=CJK_FONT, size=8.5, color="1F2937")


def _add_quote(doc: Document, lines: Sequence[str]) -> None:
    p = doc.add_paragraph(style="WP Quote")
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.right_indent = Inches(0.1)
    _shade_paragraph(p, CALLOUT_FILL)
    _set_paragraph_left_border(p, color=HEADING_BLUE)
    _add_inline(p, " ".join(line.lstrip("> ") for line in lines))


def _flush_paragraph(doc: Document, buffer: list[str]) -> None:
    if not buffer:
        return
    text = " ".join(part.strip() for part in buffer if part.strip())
    buffer.clear()
    if not text:
        return
    p = doc.add_paragraph()
    _add_inline(p, text)


LIST_RE = re.compile(r"^(\s*)([-+*]|\d+[.)])\s+(.*)$")


def add_markdown_body(doc: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    buffer: list[str] = []
    ordered_num_id: int | None = None
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()
        peek_list = LIST_RE.match(line)
        if not (peek_list and peek_list.group(2)[0].isdigit()):
            ordered_num_id = None
        if stripped.startswith("```"):
            _flush_paragraph(doc, buffer)
            language = stripped[3:].strip()
            idx += 1
            code: list[str] = []
            while idx < len(lines) and not lines[idx].strip().startswith("```"):
                code.append(lines[idx])
                idx += 1
            _add_code_block(doc, code, language)
            idx += 1
            continue
        if stripped.startswith("#"):
            match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
            if match:
                _flush_paragraph(doc, buffer)
                level = min(6, len(match.group(1)))
                p = doc.add_paragraph(style=f"Heading {level}")
                _add_inline(p, match.group(2).strip())
                idx += 1
                continue
        if stripped.startswith(">"):
            _flush_paragraph(doc, buffer)
            quote = [line]
            idx += 1
            while idx < len(lines) and lines[idx].strip().startswith(">"):
                quote.append(lines[idx])
                idx += 1
            _add_quote(doc, quote)
            continue
        if "|" in line and idx + 1 < len(lines) and _is_table_separator(lines[idx + 1]):
            _flush_paragraph(doc, buffer)
            rows = [_split_table_row(line)]
            idx += 2
            while idx < len(lines) and "|" in lines[idx] and lines[idx].strip():
                rows.append(_split_table_row(lines[idx]))
                idx += 1
            _add_table(doc, rows)
            continue
        list_match = LIST_RE.match(line)
        if list_match:
            _flush_paragraph(doc, buffer)
            indent, marker, content = list_match.groups()
            level = min(2, len(indent.expandtabs(4)) // 2)
            if marker[0].isdigit():
                if ordered_num_id is None:
                    ordered_num_id = _create_decimal_numbering(doc, start=int(re.match(r"\d+", marker).group(0)))
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.25
                _apply_numbering(p, num_id=ordered_num_id, level=level)
            else:
                base = "List Bullet"
                style = base if level == 0 else f"{base} {level + 1}"
                p = doc.add_paragraph(style=style)
            _add_inline(p, content)
            idx += 1
            continue
        if stripped in {"---", "***", "___"}:
            _flush_paragraph(doc, buffer)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            _set_paragraph_left_border(p, color="D0D5DD", size=4, space=0)
            idx += 1
            continue
        if not stripped:
            _flush_paragraph(doc, buffer)
            idx += 1
            continue
        buffer.append(line)
        idx += 1
    _flush_paragraph(doc, buffer)


def build_docx(manifest: dict, stats: dict, resolved: Sequence[dict]) -> None:
    doc = Document()
    _configure_page(doc)
    _configure_styles(doc)
    _configure_header_footer(doc)
    doc.core_properties.title = manifest["title"]
    doc.core_properties.subject = manifest["subtitle"]
    doc.core_properties.author = "Black Framework Stack Project"
    doc.core_properties.keywords = "blackbase, nsgablack, mlblack, 优化, 机器学习, 运行架构"

    _add_cover(doc, manifest, stats)
    _add_contents(doc, resolved)

    for idx, chapter in enumerate(resolved):
        if idx:
            doc.add_page_break()
        heading = doc.add_paragraph(chapter["title"], style="Heading 1")
        heading.paragraph_format.page_break_before = False
        note = doc.add_paragraph(style="WP Source Note")
        note_text = _source_note(chapter).lstrip("> ")
        _add_inline(note, note_text)
        _shade_paragraph(note, CALLOUT_FILL)
        _set_paragraph_left_border(note, color=HEADING_BLUE, size=8, space=6)
        add_markdown_body(doc, chapter["body"])

    doc.save(DOCX_OUT)


def main() -> None:
    manifest = _load_manifest()
    markdown, resolved = compose_markdown(manifest)
    stats = _stats(markdown, resolved, manifest)
    if stats["readable_characters_excluding_whitespace"] < stats["target_characters"]:
        raise RuntimeError(
            "whitepaper target not met: "
            f"{stats['readable_characters_excluding_whitespace']} < {stats['target_characters']}"
        )
    MARKDOWN_OUT.write_text(markdown, encoding="utf-8")
    STATS_OUT.write_text(json.dumps(stats, ensure_ascii=False, indent=2), encoding="utf-8")
    build_docx(manifest, stats, resolved)
    print(json.dumps({**stats, "markdown": str(MARKDOWN_OUT), "docx": str(DOCX_OUT)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
